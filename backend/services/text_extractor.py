"""
Extract text from PDF and DOCX resume files.
Kept from Phase 1 — utility only, not scoring.
"""
import io
from typing import Tuple
from docx import Document

try:
    import pdfplumber
    USE_PDFPLUMBER = True
except ImportError:
    USE_PDFPLUMBER = False


def extract_text(file_content: bytes, filename: str) -> Tuple[bool, str]:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _from_pdf(file_content)
    elif name.endswith(".docx"):
        return _from_docx(file_content)
    return False, f"Unsupported file type: {filename}"


def _from_pdf(content: bytes) -> Tuple[bool, str]:
    try:
        buf = io.BytesIO(content)
        if USE_PDFPLUMBER:
            parts = []
            with pdfplumber.open(buf) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            text = "\n".join(parts)
        else:
            from PyPDF2 import PdfReader
            reader = PdfReader(buf)
            text = "\n".join(
                p.extract_text() for p in reader.pages if p.extract_text()
            )
        if not text.strip():
            return False, "PDF appears empty or is image-based. Use a text-based PDF."
        return True, text
    except Exception as e:
        return False, f"PDF extraction failed: {str(e)}"


def _from_docx(content: bytes) -> Tuple[bool, str]:
    try:
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts)
        if not text.strip():
            return False, "DOCX appears empty."
        return True, text
    except Exception as e:
        return False, f"DOCX extraction failed: {str(e)}"
